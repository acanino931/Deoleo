# This is a sample Python script.
from src import yearly_functions as yf
import subprocess
import pandas as pd
from src import graphic_functions as gf
from docx import Document
from docx.shared import Inches
from PIL import Image
from io import BytesIO
from src import importing_data as imd
from src import Regression_Functions as rf
from src import Aux_functions as aux
import numpy as np
import statsmodels.api as sm
import os
import os
from src import Yearly_weight_transformation as year_trans
import tabula
import importlib # code to reload  lib
from unidecode import unidecode
#from src import importing_data as imd  # code to reload  lib
importlib.reload(rf)  # Reload the module # code to reload  lib

# test regression :
from datetime import datetime
from dateutil.relativedelta import relativedelta

from sklearn.linear_model import LinearRegression, HuberRegressor
from sklearn.metrics import r2_score, mean_absolute_error,mean_squared_error, mean_absolute_percentage_error
from scipy.stats import kurtosis, skew
import itertools
import statsmodels.api as sm
from statsmodels.tsa.stattools import adfuller
from statsmodels.stats.diagnostic import acorr_ljungbox
import main
from sklearn.linear_model import Ridge
from sklearn.linear_model import Lasso
from sklearn.metrics import r2_score


from sklearn.linear_model import Lasso
from sklearn.metrics import r2_score
import numpy as np
from src import Feature_selection as fs


def lasso_feature_selection(X, y, n_vars=9, alpha=1.0):
    model = Lasso(alpha=alpha)
    model.fit(X, y)

    # Select only the top 'n_vars' features based on Lasso coefficients
    top_features = np.argsort(np.abs(model.coef_))[::-1][:n_vars]

    selected_features = X.columns[top_features]
    selected_coefficients = model.coef_[top_features]

    # Predict on the training set using the selected features
    X_selected = X[selected_features]
    y_pred = model.predict(X)

    # Calculate R^2 on the training set
    r2_train = r2_score(y, y_pred)

    # Use statsmodels to calculate standard errors and p-values for the selected features
    X_selected_with_intercept = sm.add_constant(X_selected)  # Include only the selected features
    model_sm = sm.OLS(y, X_selected_with_intercept).fit()

    selected_p_values = model_sm.pvalues[1:]  # Exclude the intercept term

    return selected_features, selected_coefficients, selected_p_values, r2_train




"""
def lasso_feature_selection(X, y, n_vars=9, alpha=1.0):
    model = Lasso(alpha=alpha)
    model.fit(X, y)

    coef_abs = np.abs(model.coef_)

    coef_sorted = np.argsort(coef_abs)[::-1]
    print(coef_abs)
    if len(coef_sorted) > n_vars:
        selected_features = [X.columns[i] for i in coef_sorted[:n_vars]]
    else:
        selected_features = [X.columns[i] for i in coef_sorted]

    return selected_features
"""

def ridge_feature_selection(X, y, n_vars=4, alpha=1.0):
    model = Ridge(alpha=alpha)
    model.fit(X, y)

    coef_abs = np.abs(model.coef_)

    coef_sorted = np.argsort(coef_abs)[::-1]
    print(coef_abs)
    if len(coef_sorted) > n_vars:
        selected_features = [X.columns[i] for i in coef_sorted[:n_vars]]
    else:
        selected_features = [X.columns[i] for i in coef_sorted]

    return selected_features



def print_doc_descriptive_vars(df1,target_var ='VIRGEN_EXTRA_EUR_kg',lag_cross_corr =24,yearly_production = True):


    # THE CCF USED HERE STARTS FROM LAG 1
    df = df1.copy()

    # calculate the column year in case is not explicited
    if 'YEAR' not in df.columns:
        if 'DATE' not in df.columns:
            df['YEAR'] = df.index.year
        else:
            df['YEAR'] = df['DATE'].year
    if yearly_production == True:
        if 'HARVEST_YEAR'not in df.columns:
            df['HARVEST_YEAR'] = df['YEAR'].shift(-2)

    correlation_matrix = df.corr()


    max_correlation_with_target = correlation_matrix['VIRGEN_EXTRA_EUR_kg'].drop(['YEAR', 'VIRGEN_EXTRA_EUR_kg']).abs().sort_values(ascending=False)
    ordered_columns = max_correlation_with_target.index.tolist()
    df = df[['YEAR', 'VIRGEN_EXTRA_EUR_kg'] + ordered_columns]
    doc = Document()
    doc.add_heading('Graficas de Todas las Variables 20 10 2023', 0)
    for col in df:
        if col != target_var and col != 'YEAR' and col != 'HARVEST_YEAR':
            doc.add_paragraph('Graficas de Variable ' + col)
            # to print the yearly var change the sequent function into : scatterplot_for_years_yearly_var() scatterplot_for_years
            image_buffer = gf.scatterplot_for_years(df, col, target_var)
            doc.add_picture(image_buffer, width=Inches(3), height=Inches(2))
            doc.add_paragraph('')
            buffer_ret = gf.custom_ccf(df, col, target_var, lag_cross_corr)
            doc.add_picture(buffer_ret, width=Inches(3), height=Inches(2))
            doc.add_paragraph('')

            buff = gf.plot_and_save_variables(df, col, target_var, temp='Monthly')
            doc.add_picture(buff, width=Inches(5), height=Inches(3))
            doc.add_paragraph('')

            buffer = gf.print_correlation_over_time(df, col, target_var)
            doc.add_picture(buffer, width=Inches(5), height=Inches(3))
            doc.add_paragraph('')

    doc.save('Output/Document/EXIST_24Nov.docx')

mock = False
if mock == False:
    try:
        # Code that might raise an exception
        df_month = imd.load_data()

        # Other code that follows if no exception is raised
        print("No exception occurred")
    except ValueError  as e:
        # Handle the exception
        print(f"An exception occurred: {e}")
else:
    df_month = pd.read_excel("Output/Excel/df_month.xlsx")
    # in the import the field grossProduction_Spain_Soybean creates some problems with the scatterplot printing: investigate

if 'DATE' in df_month.columns:
    df_month  = df_month.set_index('DATE')
df_month_trans = year_trans.yearly_production_transform(path_df_month = "Output/Excel/df_month.xlsx")
if 'DATE' in df_month_trans.columns:
    df_month_trans = df_month_trans.set_index('DATE')

df_month = df_month_trans.copy()
df_month['IMPORTS_LAG_9'] = df_month['IMPORTS'].shift(9)
df_month['EXPORTS_LAG_12'] = df_month['EXPORTS'].shift(12)
df_month


df_exist = pd.read_excel("Datos/Existencias/Existencias_To_import.xlsx", sheet_name="Agosto")
#df_exist = pd.read_excel("Datos/Existencias/Existencias_To_import.xlsx", sheet_name="Agosto")

df_exist.set_index("DATE",inplace = True)

df_exist.columns
ls_exist = list(df_exist.columns)

df_month1= pd.merge(df_month, df_exist , right_index= True , left_index= True , how = "left" )

ls_exist.append('VIRGEN_EXTRA_EUR_kg')
ls_exist
df_selection = df_month1[ls_exist].copy()
df_selection = rf.eliminate_rows_from_date(df_selection, '2010-07-01')

df_res1 = fs.semimanual_single_regressions(df_selection ,  all_columns =ls_exist ,target_var = 'VIRGEN_EXTRA_EUR_kg')

df_res1



pdf = imd.include_pdf_data(df_month)
df_month = pdf


df_month.columns
#df_inflaction = pd.read_excel("Datos/inflacion_mensual.xls",sheet_name="Sheet1")
#df_inflaction.set_index('DATE',inplace = True)
#df_month = df_month.merge(df_inflaction, left_index=True, right_index=True, how='left')
#df_month
#df_inflaction.columns

df_month.columns
basic_model_2010_necessary = ['VIRGEN_EXTRA_EUR_kg', 'PRODUCTION_HARVEST', 'TOTAL_CONS','PRODUCTION_HARVEST_LAST_YEAR', 'HARVEST_FORECAST_JUNTA_ANDALUCIA']
#basic_model_2010_necessary = [ 'PRODUCTION_HARVEST', 'TOTAL_CONS','PRODUCTION_HARVEST_LAST_YEAR', 'HARVEST_FORECAST_JUNTA_ANDALUCIA']


col_energy =['VIRGEN_EXTRA_EUR_kg','Trade Close_EUA','Mid Price Close_BRENT', 'Trade Close_API2', 'Trade Close_TTF', 'Media POOL_OMEL', 'MONTHLY_INFLATION_PERC']



col_energy_necessary =['Trade Close_EUA','Mid Price Close_BRENT']

col_pdf_necessary = ['Consumo UE_TOTAL B)', 'Consumo UE_TOTAL  A)', 'Exportacion UE_TOTAL B)']# , 'Importacion UE_TOTAL B)
col_pdf_necessary = ['Consumo UE_TOTAL B)', 'Consumo UE_TOTAL  A)' ,'Consumo Total_TOTAL MONDIAL WORLD_LAG_12','Exportacion UE_TOTAL B)', 'Produccion UE_TOTAL  A)'] #  ,
col_necessary = basic_model_2010_necessary + col_energy_necessary+ col_pdf_necessary

col_necessary = basic_model_2010_necessary + col_energy_necessary
col_necessary

basic_model_df = df_month[col_necessary]
basic_model_df = rf.eliminate_rows_from_date(basic_model_df, '2010-07-01')

basic_model_df = rf.eliminate_rows_after_date(basic_model_df, '2022-09-01')

basic_model_df.columns

#df_month.drop( 'DATE.1',inplace = True, axis = 1)

df_month.columns


col_necessary


col_necessary
y= basic_model_df[['VIRGEN_EXTRA_EUR_kg']]
X = basic_model_df[col_necessary]
if 'VIRGEN_EXTRA_EUR_kg' in X.columns:
    X.drop(columns =['VIRGEN_EXTRA_EUR_kg'],inplace = True)
X = sm.add_constant(X)
model = sm.OLS(y, X).fit()
print(model.summary())


basic_model_df.columns

df_result = rf.back_testing_regression_rolling_OLD(basic_model_df, basic_model_df.drop(columns=['VIRGEN_EXTRA_EUR_kg']),
                                        'VIRGEN_EXTRA_EUR_kg', signif=False, initial_date='2019-05-01', window = 40,
                                        final_date='2023-09-01')

df_result = rf.back_testing_regression_expanding_OLD(basic_model_df, basic_model_df.drop(columns=['VIRGEN_EXTRA_EUR_kg']),
                                        'VIRGEN_EXTRA_EUR_kg', signif=False, initial_date='2019-05-01', window = 40,
                                        final_date='2023-09-01')
df_result.Mape_final.mean()


df_result.to_excel("Output/Excel/backtesting_rolling_40_basic_validated_Included_year.xlsx")

df_result.to_excel("Output/Excel/backtesting_expanding_40_basic_validated_Included_year.xlsx")


print (len(col_necessary)-1,len(X.columns)-1)
if (model.pvalues < 0.05).all(): #  and len(col_necessary) < len(X.columns)-1
    print('ok')
df_month1 = df_month[var_pdf]

df_month1 = df_month[col_necessary]
df_month1 = rf.eliminate_rows_from_date(df_month1, '2010-07-01')
basic_model_df.columns
basic_model_df.drop(columns = ['Consumo UE_TOTAL B)',
       'Consumo UE_TOTAL  A)', 'Exportacion UE_TOTAL B)_LAG_12'], inplace = True)
var_pdf
col_necessary

basic_model_df = df_month[col_necessary + var_pdf]
basic_model_df = rf.eliminate_rows_from_date(basic_model_df, '2010-07-01')
for i in col_necessary :
    if i not in df_month.columns:
        print(i)

basic_model_df.columns
basic_model_df.to_excel("Output/Excel/df_model_piu_pdf.xlsx")

df_month2 = pd.merge (basic_model_df,df_month1, left_index=True, right_index=True, how = "left")
df_month2
col_necessary
var_sig  = fs.semiautomatic_adding_feature (basic_model_df , all_columns = list(basic_model_df.columns) , model_columns =col_necessary )


var_sig


len(var_sig)

var_sig
var_sig2
var_sig2 = var_sig

df_month

basic_model_df.columns

basic_model_df = df_month[['VIRGEN_EXTRA_EUR_kg']]


basic_model_df_copy = basic_model_df.copy()
basic_model_df= basic_model_df.shift(3)
basic_model_df['VIRGEN_EXTRA_EUR_kg'] = basic_model_df['VIRGEN_EXTRA_EUR_kg'].shift(-3)

basic_model_df1 = yf.aggregate_mountly_data(basic_model_df)
basic_model_df1
basic_model_df = basic_model_df1[var_pdf_very_good].copy()
basic_model_df


basic_model_df = df_month[['VIRGEN_EXTRA_EUR_kg']]


basic_model_df.columns

basic_model_df = df_month[col_energy].copy()

basic_model_df.info()
basic_model_df['Trade Close_TTF'].loc['2010-01-01':]
basic_model_df = basic_model_df.fillna(method='ffill')
basic_model_df.columns
basic_model_df = df_month[['VIRGEN_EXTRA_EUR_kg']]

basic_model_df
pdf = imd.include_pdf_data(basic_model_df)
basic_model_df = pdf
pdf.columns

pdf


pdf_good = col_selected
pdf = pdf[pdf_good]
pdf.drop(columns=['VIRGEN_EXTRA_EUR_kg'],inplace = True)

basic_model_df= pdf



pdf.columns
basic_model_df.columns
#df_month_pdf.columns
df_andaluz = imd.include_meteo_variables()

var_pdf = ['Consumo Total_TOTAL  A',
       'Consumo Total_TOTAL B', 'Consumo Total_TOTAL MONDIAL WORLD',
       'Consumo UE_TOTAL  A)', 'Consumo UE_TOTAL B)', 'Consumo UE_TOTAL A + B',
       'Exportacion Total_TOTAL A','Exportacion Total_TOTAL B',
       'Exportacion Total_TOTAL MONDIAL WORLD', 'Exportacion UE_TOTAL  A)',
       'Exportacion UE_TOTAL B)', 'Exportacion UE_TOTAL A + B',
       'Importacion Total_TOTAL  A', 'Importacion Total_TOTAL  B',
       'Importacion Total_TOTAL MONDIAL WORLD', 'Importacion UE_TOTAL  A)',
       'Importacion UE_TOTAL B)', 'Importacion UE_TOTAL A + B',
       'Produccion Total_TOTAL  A', 'Produccion Total_TOTAL B',
       'Produccion Total_TOTAL MONDIAL WORLD', 'Produccion UE_TOTAL  A)' ,'Produccion UE_TOTAL A + B']


var_pdf_good = ['Consumo Total_TOTAL B', 'Consumo Total_TOTAL MONDIAL WORLD',
       'Consumo UE_TOTAL  A)', 'Consumo UE_TOTAL B)', 'Consumo UE_TOTAL A + B'
       'Exportacion Total_TOTAL MONDIAL WORLD', 'Exportacion UE_TOTAL  A)',
       'Exportacion UE_TOTAL B)', 'Exportacion UE_TOTAL A + B','Importacion Total_TOTAL  B',
       'Importacion Total_TOTAL MONDIAL WORLD', 'Importacion UE_TOTAL  A)',
       'Importacion UE_TOTAL B)', 'Importacion UE_TOTAL A + B',
       'Produccion Total_TOTAL  A', 'Produccion Total_TOTAL B',
       'Produccion Total_TOTAL MONDIAL WORLD','Produccion UE_TOTAL A + B']


var_pdf_very_good = [ 'VIRGEN_EXTRA_EUR_kg','Consumo UE_TOTAL A + B', 'Exportacion UE_TOTAL A + B', 'Importacion UE_TOTAL A + B',
       'Produccion Total_TOTAL  A', 'Produccion Total_TOTAL B']


var_consumo =['Consumo Total_TOTAL  A',
       'Consumo Total_TOTAL B', 'Consumo Total_TOTAL MONDIAL WORLD',
       'Consumo UE_TOTAL  A)', 'Consumo UE_TOTAL B)', 'Consumo UE_TOTAL A + B']

var_import = ['Importacion Total_TOTAL  A', 'Importacion Total_TOTAL  B',
       'Importacion Total_TOTAL MONDIAL WORLD', 'Importacion UE_TOTAL  A)',
       'Importacion UE_TOTAL B)', 'Importacion UE_TOTAL A + B']

var_produc = ['Produccion Total_TOTAL  A', 'Produccion Total_TOTAL B',
       'Produccion Total_TOTAL MONDIAL WORLD', 'Produccion UE_TOTAL  A)']

var_pdf_rolling_good =['Consumo Total_TOTAL  A', 'Exportacion Total_TOTAL B', 'Importacion Total_TOTAL MONDIAL WORLD']

var_pdf_expanding_good =['Consumo UE_TOTAL A + B', 'Produccion UE_TOTAL A + B', 'Importacion UE_TOTAL B)',
                         'Consumo UE_TOTAL  A)', 'Importacion Total_TOTAL MONDIAL WORLD']


len(var_pdf_good)

df_andaluz.columns


col_andaluz = ['cumulated_year_Precip_Average_Andalucia','cumulated_month_Precip_Average_Andalucia','TMin_Average_Andalucia']

col_andaluz_2010 = [ 'ndays_Precip_Average_Andalucia','cumulated_year_Precip_Average_Andalucia']
energy_good = ['Trade Close_EUA', 'Mid Price Close_BRENT']

basic_model_2010 = ['VIRGEN_EXTRA_EUR_kg', 'IMPORTS', 'PRODUCTION_HARVEST', 'TOTAL_CONS','PRODUCTION_HARVEST_LAST_YEAR', 'HARVEST_FORECAST_JUNTA_ANDALUCIA','IMPORTS_LAG_9', 'EXPORTS_LAG_12']
basic_model_2010


var_added = col_andaluz_2010 + energy_good

var_selected = basic_model_2010 + var_added
var_selected

df_month.columns
df_andaluz = df_andaluz[col_andaluz]
basic_model_df = df_month[var_selected]

basic_model_df.columns

df_andaluz = df_andaluz[col_andaluz_2010]

basic_model_df = df_month[['VIRGEN_EXTRA_EUR_kg']]
df_month = df_month.merge(df_andaluz, left_index=True, right_index=True, how='left')

basic_model_df = basic_model_df.merge(pdf, left_index=True, right_index=True, how='left')

basic_model_df.columns

basic_model_df_corr = basic_model_df.corr()
basic_model_df_corr.iloc[:,5:]
basic_model_df.columns

basic_model_df = df_month[['VIRGEN_EXTRA_EUR_kg']]

basic_model_df = df_month[
    ['VIRGEN_EXTRA_EUR_kg', 'EXIS_INIC', 'IMPORTS', 'EXPORTS', 'INNER_CONS',  'PRODUCTION_HARVEST',
      'TOTAL_CONS', 'PRODUCTION_HARVEST_LAST_YEAR', 'PRODUCTION_HARVEST_2_YEARS']]
# 'Estimación España (Junta Andalucia)','PRODUCTION_HARVEST_REAL_EST',]] #'INTERNAL_DEMAND','PRODUCTION',



basic_model_df['IMPORTS_LAG_9'] = df_month['IMPORTS'].shift(9)
basic_model_df['EXPORTS_LAG_12'] = df_month['EXPORTS'].shift(12)
basic_model_df['TOTAL_CONS_LAG_12'] = df_month['TOTAL_CONS'].shift(12)
basic_model_df['PRODUCTION_HARVEST_LAG_8'] = df_month['PRODUCTION_HARVEST'].shift(8)
basic_model_df['EXIS_INIC_LAG_15'] = df_month['EXIS_INIC'].shift(15)

basic_model_df.columns



basic_model_df.columns


basic_model_df.columns

# modelo benchmark 2010
basic_model_df = df_month[
    ['VIRGEN_EXTRA_EUR_kg',  'IMPORTS',  'PRODUCTION_HARVEST',
      'TOTAL_CONS', 'PRODUCTION_HARVEST_LAST_YEAR',
     'HARVEST_FORECAST_JUNTA_ANDALUCIA']] #'INTERNAL_DEMAND','PRODUCTION',
basic_model_df['IMPORTS_LAG_9'] = df_month['IMPORTS'].shift(9)
basic_model_df['EXPORTS_LAG_12'] = df_month['EXPORTS'].shift(12)

basic_model_df.columns

basic_model_df.columns
# modelo benchmark 2005 :
basic_model_df = df_month[['VIRGEN_EXTRA_EUR_kg','IMPORTS','INNER_CONS','PRODUCTION_HARVEST','TOTAL_CONS','PRODUCTION_HARVEST_LAST_YEAR','IMPORTS_LAG_9','TOTAL_CONS_LAG_12']]


#modelo sin previsiones desde el 2005:
basic_model_df = df_month[
    ['VIRGEN_EXTRA_EUR_kg', 'EXIS_INIC', 'IMPORTS', 'EXPORTS',  'PRODUCTION_HARVEST',
      'TOTAL_CONS', 'PRODUCTION_HARVEST_LAST_YEAR', 'PRODUCTION_HARVEST_2_YEARS',
     ]] #'PRODUCTION_HARVEST_REAL_EST','INTERNAL_DEMAND','PRODUCTION',

df_EXIST = df_month[['VIRGEN_EXTRA_EUR_kg','EXIS_INIC']]

print_doc_descriptive_vars(basic_model_df, target_var='VIRGEN_EXTRA_EUR_kg', lag_cross_corr=24)
# END  # graficas Modelo basico Review

# SECOND VERSION Modelo Basico

df_month['IMPORTS_LAG_9'] = df_month['IMPORTS'].shift(9)
df_month['EXPORTS_LAG_12'] = df_month['EXPORTS'].shift(12)
df_month['TOTAL_CONS_LAG_12'] = df_month['TOTAL_CONS'].shift(12)
df_month['PRODUCTION_HARVEST_LAG_8'] = df_month['PRODUCTION_HARVEST'].shift(8)

df_month['PRODUCTION_HARVEST_REAL_EST_LAG_14'] = df_month['PRODUCTION_HARVEST_REAL_EST'].shift(14)
df_month['EXIS_INIC_LAG_15'] = df_month['EXIS_INIC'].shift(15)
#basic_model_df['PRODUCTION_18'] = basic_model_df['PRODUCTION'].shift(18)
# this is probably a bad var basic_model_df['INTERNAL_DEMAND_12'] = basic_model_df['INTERNAL_DEMAND'].shift(12)
#basic_model_df['DP_PRODUCTION_HARVEST_LAG_14'] = basic_model_df['DP_PRODUCTION_HARVEST'].shift(14)
#basic_model_df = rf.eliminate_rows_from_date(basic_model_df, '2005-10-01')

basic_model_df.columns

basic_model_df

#print_doc_descriptive_vars(basic_model_df)
basic_model_df = rf.eliminate_rows_from_date(basic_model_df, '2005-10-01')
basic_model_df


basic_model_df = rf.eliminate_rows_from_date(basic_model_df, '2010-07-01')

basic_model_df

basic_model_df

df_month_copy = rf.eliminate_rows_from_date(df_month_copy, '2010-07-01')
df_month_copy

basic_model_df.index

basic_model_df

basic_model_df = basic_model_df[basic_model_df.index > 2009]
basic_model_df= basic_model_df.fillna(method='ffill')

df_month_copy = df_month[basic_model_2010]
basic_model_2010


df_month_copy.columns

basic_model_df.info()

target_variable = 'VIRGEN_EXTRA_EUR_kg'
y = basic_model_df[[target_variable]].copy()
X = basic_model_df.drop(columns=[target_variable ]).copy()

rows_with_null = X[X.isnull().any(axis=1)]
rows_with_null


basic_model_df.columns

basic_model_df.HARVEST_FORECAST_JUNTA_ANDALUCIA
basic_model_df.columns
basic_model_df.drop('Importacion UE_TOTAL B)',axis = 1, inplace = True)


target_variable = 'VIRGEN_EXTRA_EUR_kg'
y = basic_model_df[[target_variable]].copy()
X = basic_model_df.drop(columns=[target_variable ]).copy()
#X = X.drop(columns=['INNER_CONS' ]).copy()
#X = basic_model_df_man[['Consumo Total_TOTAL  A','Consumo Total_TOTAL B']]
X = sm.add_constant(X)
model = sm.OLS(y, X).fit()
print(model.summary())


X.columns
basic_model_df.columns
col_drop = ['Produccion Total_TOTAL B', 'Importacion UE_TOTAL A + B' ,'EXPORTS_LAG_12','Produccion Total_TOTAL  A','Produccion Total_TOTAL MONDIAL WORLD',
            'Consumo Total_TOTAL B','Consumo Total_TOTAL  A','Exportacion UE_TOTAL  A)','Exportacion UE_TOTAL  A)']
basic_model_dfman= basic_model_df.drop(columns=col_drop)
target_variable = 'VIRGEN_EXTRA_EUR_kg'
y = basic_model_dfman[[target_variable]].copy()
X = basic_model_dfman.drop(columns=[target_variable ]).copy()
#X = basic_model_df_man[['Consumo Total_TOTAL  A','Consumo Total_TOTAL B']]

X = sm.add_constant(X)
model = sm.OLS(y, X).fit()
print(model.summary())

basic_model_df
# start stepwise
len(basic_model_df.columns)
basic_model_step = basic_model_df.copy()
#basic_model_step = rf.eliminate_rows_from_date(basic_model_df, '2005-10-01')
#basic_model_step =basic_model_dfman
iteration_selected = 6
df_step = rf.stepwise_eliminating(basic_model_step, 'VIRGEN_EXTRA_EUR_kg', iteration_selected)
df_step.columns
print(df_step.iloc[:, 0:4])

rf.save_model_summary_to_file(df_step, iteration_selected)
col_selected = df_step.loc[df_step.index[iteration_selected - 1], 'Actual_cols']

len(col_selected)
col_selected

gf.plot_correlation_target_variable(df_month_copy, target_variable = 'VIRGEN_EXTRA_EUR_kg')

basic_model_df.columns

target_variable = 'VIRGEN_EXTRA_EUR_kg'
y = basic_model_df[[target_variable]].copy()
X = basic_model_df[col_selected].copy()
X = X.drop(columns=[target_variable ]).copy()
#X = basic_model_df_man[['Consumo Total_TOTAL  A','Consumo Total_TOTAL B']]
X = sm.add_constant(X)
model = sm.OLS(y, X).fit()
print(model.summary())
X.columns

basic_model_df.columns

col_necessary
basic_model_df = basic_model_df[col_necessary]

basic_model_df

basic_model_df.columns
ls_ventanas = [12,24, 30, 36, 48, 50, 60 ]
for window in ls_ventanas:
    result_rolling = rf.rolling_regression(basic_model_df, 'VIRGEN_EXTRA_EUR_kg', window)
    result_rolling = aux.add_average_row(result_rolling)
    result_rolling.to_excel(f"Output/Rolling_Regression/Rolling{window}.xlsx")


# rolling stepwise
#col_selected.append('VIRGEN_EXTRA_EUR_kg')
col_selected
basic_model_df_roll = basic_model_df[col_selected]
basic_model_df_roll.columns

for window in ls_ventanas:
    result_rolling = rf.rolling_regression(basic_model_df_roll, 'VIRGEN_EXTRA_EUR_kg', window)
    result_rolling = aux.add_average_row(result_rolling)
    result_rolling.to_excel(f"Output/Rolling_Regression/Rolling{window}.xlsx")


basic_model_df_bck = basic_model_step[col_selected].copy()
df_step.loc[iteration_selected - 1, 'Model_summary']


len(X.columns)
# lasso
selected_features, selected_coefficients, selected_p_values, r2_train = lasso_feature_selection(X, y, n_vars=9, alpha=1.0)

print("Selected Features:", selected_features)
print("Corresponding Coefficients:", selected_coefficients)
print("P-values:", selected_p_values)
print(f"R^2 on Training Set: {r2_train}")



basic_model_df = df_month[col_necessary]

# Granger :

import numpy as np
import pandas as pd
from statsmodels.tsa.stattools import grangercausalitytests
import matplotlib.pyplot as plt

basic_model_df.columns

df_granger = basic_model_df[['TOTAL_CONS','VIRGEN_EXTRA_EUR_kg']]

df_granger = df_month[['EXPORTS','VIRGEN_EXTRA_EUR_kg']]

df_granger = df_month[['EXPORTS','VIRGEN_EXTRA_EUR_kg']]

df_granger = df_month[['VIRGEN_EXTRA_EUR_kg','INNER_CONS']]

df_reverse_granger = basic_model_df[['TOTAL_CONS','VIRGEN_EXTRA_EUR_kg']]
# Generate some example data
df_granger = rf.eliminate_rows_from_date(df_granger, '2010-07-01')
df_reverse_granger = rf.eliminate_rows_from_date(df_reverse_granger, '2010-07-01')


# testing if total cons explains y
# Create a DataFrame

df_granger.info()

# Plot the data
df_granger.plot(title='Oil and Olive Oil Prices Over Time')
plt.show()

# Perform Granger causality test
max_lag = 13  # you can adjust this based on your data
test_result = grangercausalitytests(df_granger, max_lag, verbose=True)

test_result

# Print the results
for lag in range(1, max_lag + 1):
    p_value = test_result[lag][0]['ssr_ftest'][1]
    print(f'Granger causality test (lag={lag}): p-value = {p_value}')


# reverse
test_result_reverse = grangercausalitytests(df_reverse_granger, max_lag, verbose=True)

# Print the results
for lag in range(1, max_lag + 1):
    p_value = test_result_reverse[lag][0]['ssr_ftest'][1]
    print(f'Granger causality test (lag={lag}): p-value = {p_value}')