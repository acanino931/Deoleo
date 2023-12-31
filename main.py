# This is a sample Python script.
from src import yearly_functions as yf
import subprocess
import pandas as pd
from src import graphic_functions as gf
from docx import Document
from docx.shared import Inches
from PIL import Image
from io import BytesIO
from src import importing_data as imd
from src import Regression_Functions as rf
from src import Aux_functions as aux
import numpy as np
import statsmodels.api as sm
import os
from src import Yearly_weight_transformation as year_trans


import tabula
import importlib # code to reload  lib
from unidecode import unidecode
#from src import importing_data as imd  # code to reload  lib
importlib.reload(rf)  # Reload the module # code to reload  lib

# test regression :
from datetime import datetime
from dateutil.relativedelta import relativedelta

from sklearn.linear_model import LinearRegression, HuberRegressor
from sklearn.metrics import r2_score, mean_absolute_error,mean_squared_error, mean_absolute_percentage_error
from scipy.stats import kurtosis, skew
import itertools
import statsmodels.api as sm
from statsmodels.tsa.stattools import adfuller
from statsmodels.stats.diagnostic import acorr_ljungbox



# Press Shift+F10 to execute it or replace it with your code.
# Press Double Shift to search everywhere for classes, files, tool windows, actions, and settings.




def print_doc_descriptive_vars(df1,target_var ='VIRGEN_EXTRA_EUR_kg',lag_cross_corr =24,yearly_production = True):
    # THE CCF USED HERE STARTS FROM LAG 1
    df = df1.copy()

    # calculate the column year in case is not explicited
    if 'YEAR' not in df.columns:
        if 'DATE' not in df.columns:
            df['YEAR'] = df.index.year
        else:
            df['YEAR'] = df['DATE'].year
    if yearly_production == True:
        if 'HARVEST_YEAR'not in df.columns:
            df['HARVEST_YEAR'] = df['YEAR'].shift(-2)

    correlation_matrix = df.corr()


    max_correlation_with_target = correlation_matrix['VIRGEN_EXTRA_EUR_kg'].drop(['YEAR', 'VIRGEN_EXTRA_EUR_kg']).abs().sort_values(ascending=False)
    ordered_columns = max_correlation_with_target.index.tolist()
    df = df[['YEAR', 'VIRGEN_EXTRA_EUR_kg'] + ordered_columns]
    doc = Document()
    doc.add_heading('Graficas de Todas las Variables 20 10 2023', 0)
    for col in df:
        if col != target_var and col != 'YEAR' and col != 'HARVEST_YEAR':
            doc.add_paragraph('Graficas de Variable ' + col)
            # to print the yearly var change the sequent function into : scatterplot_for_years_yearly_var() scatterplot_for_years
            image_buffer = gf.scatterplot_for_years(df, col, target_var)
            doc.add_picture(image_buffer, width=Inches(3), height=Inches(2))
            doc.add_paragraph('')
            buffer_ret = gf.custom_ccf(df, col, target_var, lag_cross_corr)
            doc.add_picture(buffer_ret, width=Inches(3), height=Inches(2))
            doc.add_paragraph('')

            buff = gf.plot_and_save_variables(df, col, target_var, temp='Monthly')
            doc.add_picture(buff, width=Inches(5), height=Inches(3))
            doc.add_paragraph('')

            buffer = gf.print_correlation_over_time(df, col, target_var)
            doc.add_picture(buffer, width=Inches(5), height=Inches(3))
            doc.add_paragraph('')

    doc.save('Output/Document/sample_with_pycharm.docx')


def print_doc_scatter_ouliers(df1,target_var ='VIRGEN_EXTRA_EUR_kg'):

    df = df1.copy()

    # calculate the column year in case is not explicited
    if 'YEAR' not in df.columns:
        if 'DATE' not in df.columns:
            df['YEAR'] = df.index.year
        else:
            df['YEAR'] = df['DATE'].year
    correlation_matrix = df.corr()


    max_correlation_with_target = correlation_matrix[target_var].drop(['YEAR', target_var]).abs().sort_values(ascending=False)
    ordered_columns = max_correlation_with_target.index.tolist()
    df = df[['YEAR', target_var] + ordered_columns]
    doc = Document()
    doc.add_heading('Graficas de Todas las Variables 20 10 2023', 0)
    df_outlier = gf.mark_outliers(df, target_var)
    outlier_column = df_outlier.columns[1]
    outlier_count = df_outlier[outlier_column].sum()
    if outlier_count > 0:
        img = gf.plot_time_series_with_outliers(df_outlier, target_var, target_var +'_is_outlier')
        doc.add_paragraph('Graficas de Outliers de la variable objetivo: ' + target_var)
        doc.add_picture(img, width=Inches(4), height=Inches(2.5))
        doc.add_paragraph('')
    for col in df:
        if col != target_var and col != 'YEAR':
            doc.add_paragraph('Graficas de Variable ' + col)
            image_buffer = gf.scatterplot_for_years(df, col, target_var)
            doc.add_picture(image_buffer, width=Inches(3), height=Inches(2))
            doc.add_paragraph('')
            df_outlier = gf.mark_outliers(df, col)
            outlier_column = df_outlier.columns[1]
            outlier_count = df_outlier[outlier_column].sum()
            if outlier_count > 0:
                img = gf.plot_time_series_with_outliers(df_outlier, col, col + '_is_outlier')
                doc.add_picture(img, width=Inches(4), height=Inches(2.5))
                doc.add_paragraph('')
            buff = gf.plot_and_save_variables(df, col, target_var, temp='Monthly')
            doc.add_picture(buff, width=Inches(5), height=Inches(3))
            doc.add_paragraph('')


    doc.save('Output/Document/Scatterplots_Modelo_basico.docx')

def include_meteo_variables(df):
    df_cordoba = imd.import_meteo_single_province("Datos/Datos_Cordoba", "Cordoba")
    df_jaen = imd.import_meteo_single_province("Datos/Datos_Jaen", "Jaen")
    df_meteo = df_jaen.merge(df_cordoba, left_index=True, right_index=True, how='left')
    df = df.merge(df_meteo, left_index=True, right_index=True, how='left')
    return df

# Press the green button in the gutter to run the script.
if __name__ == '__main__':

    mock = False
    if mock == False:
        try:
            # Code that might raise an exception
            df_month = imd.load_data()

            # Other code that follows if no exception is raised
            print("No exception occurred")
        except ValueError  as e:
            # Handle the exception
            print(f"An exception occurred: {e}")
    else:
        df_month = pd.read_excel("Output/Excel/df_month.xlsx")
        # in the import the field grossProduction_Spain_Soybean creates some problems with the scatterplot printing: investigate

    if 'DATE' in df_month.columns:
        df_month  = df_month.set_index('DATE')
        df_month.columns
    df_month_trans = year_trans.yearly_production_transform(path_df_month = "Output/Excel/df_month.xlsx")
    if 'DATE' in df_month_trans.columns:
        df_month_trans = df_month_trans.set_index('DATE')

    df_month = df_month_trans.copy()

    df = df_month_trans[['VIRGEN_EXTRA_EUR_kg','PRODUCTION_HARVEST_REAL_EST','PRODUCTION_HARVEST','DP_PRODUCTION_HARVEST']].copy()
    df_month = rf.eliminate_rows_from_date(df, '2005-10-01')

    #df_month_pdf.to_excel("Output/Excel/Fisicas_pdf_2nd_Review.xlsx")

    df_month = df_month_trans.copy()
   # df_month = df_month[['VIRGEN_EXTRA_EUR_kg']].copy()
    df_month = include_meteo_variables(df_month)
    #df_pdf = imd.import_pdf_data()
    df_month_pdf = imd.include_pdf_data(df_month)
    df_month_pdf.columns
    df_month_pdf = rf.eliminate_rows_from_date(df_month_pdf, '2005-10-01')

    df_month_pdf['Produccion UE_TOTAL A + B'].dtype
    df_month_pdf.to_excel("Prova.xlsx")



    df_month.columns
    #df_import_pdf_data()
    #df_test = basic_model_df.copy()
    print_doc_descriptive_vars(df_month)
    basic_model_df


    basic_model_df_man
    ls_ventanas= [24,30,36,40,50]
    for window in ls_ventanas:
        result_rolling = rf.rolling_regression(basic_model_df_man, 'VIRGEN_EXTRA_EUR_kg', 40)
        result_rolling= aux.add_average_row(result_rolling)
        result_rolling.to_excel(f"Output/Rolling_Regression/Rolling{window}.xlsx")


    result_rolling.to_excel("Prova_rolling.xlsx")


    # analisis harvest production
    #   df_test = rf.eliminate_rows_after_date(df_test, '2021-12-01')
    #   df_test = rf.eliminate_rows_from_date(df_test, '2005-10-01')



    #df_test = df_month[['VIRGEN_EXTRA_EUR_kg', 'VIRGEN_EUR_kg','HARVEST_YEAR','PRODUCTION_HARVEST','YEAR']]
    print_doc_descriptive_vars(df_test, target_var='VIRGEN_EXTRA_EUR_kg', lag_cross_corr=24)
    print_doc_scatter_ouliers(df_test, target_var='VIRGEN_EXTRA_EUR_kg')

    # graficas Modelo basico Review :
    df_month = df_month.fillna(method='ffill')
    df_month.columns
    df_month = df_month[['VIRGEN_EXTRA_EUR_kg', 'EXIS_INIC', 'IMPORTS', 'EXPORTS', 'INNER_CONS', 'PRODUCTION', 'PRODUCTION_HARVEST','INTERNAL_DEMAND','TOTAL_DEMAND','TOTAL_CONS']].copy()
    df_month = df_month[['VIRGEN_EXTRA_EUR_kg', 'EXIS_INIC', 'IMPORTS', 'EXPORTS', 'INNER_CONS', 'PRODUCTION', 'PRODUCTION_HARVEST','INTERNAL_DEMAND','TOTAL_DEMAND','TOTAL_CONS',
                               'PRODUCTION_HARVEST_REAL_EST', 'PRODUCTION_HARVEST_LAST_YEAR','PRODUCTION_HARVEST_2_YEARS','DP_PRODUCTION_HARVEST' ]]

    basic_model_df = df_month[['VIRGEN_EXTRA_EUR_kg', 'EXIS_INIC', 'IMPORTS', 'EXPORTS', 'INNER_CONS', 'PRODUCTION', 'PRODUCTION_HARVEST','INTERNAL_DEMAND','TOTAL_DEMAND','TOTAL_CONS',
                               'PRODUCTION_HARVEST_REAL_EST', 'PRODUCTION_HARVEST_LAST_YEAR','PRODUCTION_HARVEST_2_YEARS','DP_PRODUCTION_HARVEST' ]]
   # basic_model_df.drop(columns=['EXTERNAL_DEMAND'], inplace=True)

    #print_doc_descriptive_vars(basic_model_df, target_var='VIRGEN_EXTRA_EUR_kg', lag_cross_corr=24)
    #END  # graficas Modelo basico Review

    # SECOND VERSION Modelo Basico

    basic_model_df['IMPORTS_LAG_21'] = basic_model_df['IMPORTS'].shift(21)
    basic_model_df['TOTAL_CONS_LAG_12'] = basic_model_df['TOTAL_CONS'].shift(12)
    basic_model_df['PRODUCTION_HARVEST_LAG_8'] = basic_model_df['PRODUCTION_HARVEST'].shift(8)
    basic_model_df['EXPORTS_LAG_12'] = basic_model_df['EXPORTS'].shift(12)
    basic_model_df['TOTAL_DEMAND_LAG_12'] = basic_model_df['TOTAL_DEMAND'].shift(12)
    basic_model_df['PRODUCTION_18'] = basic_model_df['PRODUCTION'].shift(18)
    basic_model_df['INTERNAL_DEMAND_12'] = basic_model_df['INTERNAL_DEMAND'].shift(12)
    basic_model_df['PRODUCTION_HARVEST_REAL_EST_LAG_14'] = basic_model_df['PRODUCTION_HARVEST_REAL_EST'].shift(14)
    basic_model_df['DP_PRODUCTION_HARVEST_LAG_14'] = basic_model_df['DP_PRODUCTION_HARVEST'].shift(14)
    basic_model_df = rf.eliminate_rows_from_date(basic_model_df, '2005-10-01')

    basic_model_df



    # Modelo Basico

    # pretreat variables :
    #'PRODUCTION_HARVEST_REAL_EST'
    #PRODUCTION_HARVEST_REAL_EST
    df_month = df_month.fillna(method='ffill')
    basic_model_df = df_month[['VIRGEN_EXTRA_EUR_kg', 'EXIS_INIC', 'IMPORTS', 'EXPORTS', 'INNER_CONS', 'PRODUCTION' ,'PRODUCTION_HARVEST', 'PRODUCTION_HARVEST_LAST_YEAR', 'PRODUCTION_HARVEST_2_YEARS', 'INTERNAL_DEMAND', 'TOTAL_CONS']].copy()
    basic_model_df = df_month.copy()
    basic_model_df['EXPORTS_LAG15'] = basic_model_df['EXPORTS'].shift(15)
    basic_model_df['INTERNAL_DEMAND_LAG_13'] = basic_model_df['INTERNAL_DEMAND'].shift(13)
    basic_model_df['TOTAL_CONS_LAG_12'] = basic_model_df['TOTAL_CONS'].shift(12)
    basic_model_df['TOTAL_CONS_LAG_13'] = basic_model_df['TOTAL_CONS'].shift(12)
    basic_model_df['PRODUCTION_LAG_21'] = basic_model_df['PRODUCTION'].shift(21)
    basic_model_df['TOTAL_CONS_LAG_13'] = basic_model_df['TOTAL_CONS'].shift(13)
    basic_model_df['EXIS_INIC_18'] = basic_model_df['EXIS_INIC'].shift(18)
    basic_model_df['PRODUCTION_LAG_21'] = basic_model_df['PRODUCTION'].shift(21)
    #basic_model_df.drop(columns=['EXTERNAL_DEMAND'], inplace=True)

    basic_model_df_man = basic_model_df[
        ['VIRGEN_EXTRA_EUR_kg', 'IMPORTS', 'INNER_CONS', 'TOTAL_CONS_LAG_12', 'EXPORTS_LAG15', 'TOTAL_CONS',
'INTERNAL_DEMAND_LAG_13', 'EXIS_INIC', 'PRODUCTION_HARVEST_LAG_8', 'PRODUCTION', 'INTERNAL_DEMAND']]

    basic_model_df_man = basic_model_df[
        ['VIRGEN_EXTRA_EUR_kg', 'IMPORTS','EXPORTS', 'INNER_CONS', 'TOTAL_CONS_LAG_12', 'EXPORTS_LAG15', 'TOTAL_CONS',
         'INTERNAL_DEMAND_LAG_13', 'EXIS_INIC','PRODUCTION_HARVEST', 'PRODUCTION_HARVEST_LAST_YEAR', 'PRODUCTION_HARVEST_2_YEARS','PRODUCTION', 'INTERNAL_DEMAND','PRODUCTION_LAG_21','TOTAL_CONS_LAG_13']]
    #modelo basico without lag
    #basic_model_df_man = basic_model_df[['VIRGEN_EXTRA_EUR_kg', 'IMPORTS', 'INNER_CONS', 'TOTAL_CONS','EXIS_INIC', 'PRODUCTION_HARVEST', 'INTERNAL_DEMAND']]
    basic_model_df_man = rf.eliminate_rows_from_date(basic_model_df_man, '2005-10-01')



    print_doc_scatter_ouliers(basic_model_df_man)

    basic_model_df.to_excel("Output/Excel/basic_model_df.xlsx")

    gf.plot_correlation_target_variable(basic_model_df, 'IMC_EXTRA_VIRGEN_EUR_kg')

    gf.plot_correlation_matrix(basic_model_df)

    basic_model_df.columns
    # manual, stepwise omitted
    basic_model_df_man = basic_model_df[
        ['VIRGEN_EXTRA_EUR_kg', 'IMPORTS', 'INNER_CONS', 'TOTAL_CONS_LAG_12', 'EXPORTS_LAG15', 'TOTAL_CONS',
         'INTERNAL_DEMAND_LAG_13', 'EXIS_INIC', 'PRODUCTION_HARVEST_LAG_8', 'PRODUCTION', 'INTERNAL_DEMAND']]

    x_variable = 'VIRGEN_EXTRA_EUR_kg'
    target_variable= 'VIRGEN_EXTRA_EUR_kg'
    basic_model_df_man = rf.eliminate_rows_from_date(basic_model_df_man, '2005-10-01')
    #basic_model_df = rf.eliminate_rows_from_date(basic_model_df, '2005-10-01')

    #original version
    basic_model_df_man = basic_model_df[['VIRGEN_EXTRA_EUR_kg', 'EXIS_INIC', 'IMPORTS', 'EXPORTS', 'INNER_CONS','PRODUCTION', 'PRODUCTION_HARVEST', 'INTERNAL_DEMAND', 'TOTAL_DEMAND','TOTAL_CONS', 'IMPORTS_LAG_21', 'TOTAL_CONS_LAG_21','PRODUCTION_HARVEST_LAG_8', \
                        'EXPORTS_LAG_12', 'TOTAL_DEMAND_LAG_12','PRODUCTION_18', 'INTERNAL_DEMAND_12']]

    basic_model_df_man = basic_model_df[['VIRGEN_EXTRA_EUR_kg', 'EXIS_INIC', 'IMPORTS', 'INNER_CONS','PRODUCTION', 'PRODUCTION_HARVEST','TOTAL_CONS', 'PRODUCTION_HARVEST_LAG_8', \
                       'PRODUCTION_18', 'INTERNAL_DEMAND_12']]

    # comes from stepwise
    basic_model_df_man = basic_model_df_man[col_selected]
    basic_model_df_man.columns
    basic_model_df_man= basic_model_df_man.drop(columns=['IMPORTS_LAG_21']).copy()

    basic_model_df_man = basic_model_df_bck.copy()
    basic_model_df_man.columns
    target_variable = 'VIRGEN_EXTRA_EUR_kg'
    y = basic_model_df_man[[target_variable]].copy()
    X = basic_model_df_man.drop(columns=[target_variable]).copy()
    X = sm.add_constant(X)
    model = sm.OLS(y, X).fit()
    print(model.summary())

    #original
    gf.plot_correlation_matrix(basic_model_df_man)
    # step
    gf.plot_correlation_matrix(basic_model_df_man[col_selected])

    # end modelo basico manual

    # start stepwise

    basic_model_step = basic_model_df.copy()
    basic_model_step = rf.eliminate_rows_from_date(basic_model_df, '2005-10-01')
    iteration_selected = 1
    df_step = rf.stepwise_eliminating(basic_model_step, 'VIRGEN_EXTRA_EUR_kg', iteration_selected)
    df_step.columns
    print(df_step.iloc[:, 0:4])

    rf.save_model_summary_to_file(df_step, iteration_selected)
    col_selected = df_step.loc[df_step.index[iteration_selected - 1], 'Actual_cols']

    len(col_selected)
    basic_model_df_bck = basic_model_step[col_selected].copy()
    df_step.loc[iteration_selected-1, 'Model_summary']
   # len(basic_model_df_bck['VIRGEN_EXTRA_EUR_kg'])
    #  df_pred, MSFE,MAPE = rf.back_testing_actual_time(basic_model_df_bck,50, 24, 'VIRGEN_EXTRA_EUR_kg') # montly model 50 obs out # 24 horizons previewd
    # df_pred.columns
    basic_model_df_bck.columns

    df_result=  back_testing_regression_OLD(basic_model_df, basic_model_df.drop(columns=['VIRGEN_EXTRA_EUR_kg']), 'VIRGEN_EXTRA_EUR_kg',  signif= True, initial_date = '2022-05-01',  final_date= '2023-05-01')
    df_result.to_excel("Prova.xlsx")

    df_pred, MSFE, MAPE = rf.back_testing_regression(basic_model_df_bck, 50, 24,
                                                     'VIRGEN_EXTRA_EUR_kg')  # montly model 50 obs out # 24 horizons previewd

    MAPE

    # end stepwise

    # model imc start
    basic_model_df = df_month.copy()

    basic_model_df.columns
    basic_model_df = gf.compute_lags_for_custom_ccf_IMC(basic_model_df, 6 )
    basic_model_df = rf.eliminate_rows_from_date(basic_model_df, '2018-01-01')
    corr_crossimc = gf.custom_ccf_IMC(basic_model_df, 'IMC_EXTRA_VIRGEN_EUR_kg')
    basic_model_df.columns

    basic_model_df.to_excel("Output/Excel/df_IMC.xlsx")
    basic_model_df.columns

    df_month.columns
    df_month1 = df_month[['VIRGEN_EXTRA_EUR_kg','Month','YEAR','SUNFLOWER_OIL','IMC_EXTRA_VIRGEN_EUR_kg']].copy()
    df_month1 = rf.eliminate_rows_from_date(df_month1, '2018-01-01')
    print_doc_descriptive_vars(df_month1, target_var='IMC_EXTRA_VIRGEN_EUR_kg', lag_cross_corr=6)

    out = gf.cross_correlation_variable_out_df(basic_model_df, 'VIRGEN_EXTRA_EUR_kg', 'IMC_EXTRA_VIRGEN_EUR_kg', 5)

    out_custom,df = gf.ccustom_ccf(basic_model_df, 'VIRGEN_EXTRA_EUR_kg','IMC_EXTRA_VIRGEN_EUR_kg' ,24)


    df.iloc[:,-2:]

    # imc model end



    #selected_datetime = '2021-12-01'
    #df_month_copy= df_month.loc[:selected_datetime].copy()
    df_month_copy


    #def compute_montly_weight(df_month,'VIRGEN_EXTRA_EUR_kg')
    df_selected = yf.select_annual_variable_from_dic (df_month_copy)
    df_aggr = yf.aggregate_mountly_data(df_selected)
    df_selected.tail(14)
    df_selected.columns
    #df_month_copy = df_month.copy()
    df_aggr.columns
    target_transposed = yf.transpose_target_variable(df_month_copy, 'VIRGEN_EXTRA_EUR_kg')
    df_aggr.index

    df_weight= yf.calc_correlations_yearly(df_aggr, target_transposed)
    df_weight.to_excel("Output/Excel/df_weight.xlsx")



    #df_month.columns = df_month.columns.str.replace('_seed', '')



# 1st basic model
    # consider to add stock oil
    #basic_model_df =  df_month[['VIRGEN_EXTRA_EUR_kg','EXIS_INIC','IMPORTS','EXPORTS', 'PRODUCTION','PRODUCTION_HARVEST','INTERNAL_DEMAND', 'EXTERNAL_DEMAND']].copy()




    gf.plot_correlation_target_variable(basic_model_df,'VIRGEN_EXTRA_EUR_kg')

#    basic_model_df =  df_month[['VIRGEN_EXTRA_EUR_kg','EXIS_INIC','IMPORTS','EXPORTS', 'PRODUCTION','PRODUCTION_HARVEST']].copy()
 #   basic_model_df =  df_month[['VIRGEN_EXTRA_EUR_kg','IMPORTS','EXPORTS']].copy()


    df_pred
    MAPE
    basic_model_df_bck.columns



    basic_model_df.columns

    selected_rows
    selected_rows1
    df_month_copy.Month
    print_doc_descriptive_vars(selected_rows,'VIRGEN_EXTRA_EUR_kg',24)
    print_doc_descriptive_vars(selected_rows1, 'VIRGEN_EXTRA_EUR_kg', 24)
    gf.plot_correlation_matrix(basic_model_df_man)

    df_month.columns
    basic_model_df_man
#   anticipo -2 production harvest -1 exports (intorno ai 40 y 35)

    basic_model_df_man = df_month[['VIRGEN_EXTRA_EUR_kg', 'EXIS_INIC', 'IMPORTS', 'EXPORTS', 'PRODUCTION','INNER_CONS' ,'PRODUCTION_HARVEST','INTERNAL_DEMAND', 'EXTERNAL_DEMAND', ]].copy()
    basic_model_df_man['PRODUCTION_HARVEST_OLD'] = basic_model_df_man['PRODUCTION_HARVEST'].shift(12)
    basic_model_df_man['PRODUCTION_HARVEST_OLD'] = basic_model_df_man['PRODUCTION_HARVEST_OLD'].fillna(method='ffill' ,limit=12)
    basic_model_df_man['PRODUCTION_HARVEST'] = basic_model_df_man['PRODUCTION_HARVEST'].shift(-3)
    basic_model_df_man['PRODUCTION_HARVEST'] = basic_model_df_man['PRODUCTION_HARVEST'].fillna(method='ffill' ,limit=3)
    basic_model_df_man['EXPORTS'] = basic_model_df_man['EXPORTS'].shift(12)
    basic_model_df_man['EXPORTS'] = basic_model_df_man['EXPORTS'].fillna(method='ffill' ,limit=12)
    basic_model_df_man['INTERNAL_DEMAND'] = basic_model_df_man['INTERNAL_DEMAND'].shift(13)
    basic_model_df_man['INTERNAL_DEMAND'] = basic_model_df_man['INTERNAL_DEMAND'].fillna(method='ffill' ,limit=13)
    basic_model_df_man['EXTERNAL_DEMAND'] = basic_model_df_man['EXTERNAL_DEMAND'].shift(13)
    basic_model_df_man['EXTERNAL_DEMAND'] = basic_model_df_man['EXTERNAL_DEMAND'].fillna(method='ffill' ,limit=13)
    basic_model_df_man['PRODUCTION'] = basic_model_df_man['PRODUCTION'].shift(6)
    basic_model_df_man['PRODUCTION'] = basic_model_df_man['PRODUCTION'].fillna(method='ffill' ,limit=6)


   #    basic_model_df_man = df_month[['VIRGEN_EXTRA_EUR_kg', 'EXIS_INIC', 'IMPORTS', 'EXPORTS', 'PRODUCTION','INNER_CONS' ,'PRODUCTION_HARVEST','INTERNAL_DEMAND', 'EXTERNAL_DEMAND',PRODUCTION_HARVEST_OLD ]].copy()
    basic_model_df_man = basic_model_df_man[['VIRGEN_EXTRA_EUR_kg', 'IMPORTS', 'INNER_CONS' , 'EXPORTS', 'PRODUCTION_HARVEST','EXTERNAL_DEMAND','PRODUCTION_HARVEST_OLD','INTERNAL_DEMAND']].copy()
    basic_model_df_man = basic_model_df_man.drop(columns = ['EXTERNAL_DEMAND','PRODUCTION'])
    basic_model_df_man = sm.add_constant(basic_model_df_man)
    target_variable = 'VIRGEN_EXTRA_EUR_kg'
    basic_model_df_man = rf.eliminate_rows_from_date(basic_model_df_man, '2005-10-01')
    y = basic_model_df_man[target_variable].copy()
    X = basic_model_df_man.drop(columns = [target_variable])
    model = sm.OLS(y, X).fit()
    print (model.summary())
    df_pred, MSFE, MAPE = rf.back_testing_regression(basic_model_df_man, 50, 24 ,'VIRGEN_EXTRA_EUR_kg')  # montly model 50 obs out # 24 horizons previewd
    MAPE
    model_sar = SARIMAX(y, exog=X, order=(0, 0, 0), seasonal_order=(0, 0, 0, 0))
    model_fit = model_sar.fit()
    print(model_fit.summary())
    MAPE_df = pd.DataFrame(MAPE)
    MSFE_df = pd.DataFrame(MSFE)
    df_mapes = pd.merge(MAPE_df,MSFE_df , left_index = True, right_index= True  )
    df_mapes.to_excel("Output/Excel/MAPE.xlsx")
    type(MAPE)
    print(X[X.isnull()])

    X.info()




    #def back_testing_actual_time(X, y, horizontes):



    """
    X, y, max_data_no_missing, column_data_max = rf.remove_null_rows(basic_model_df,target_variable = 'VIRGEN_EXTRA_EUR_kg')

    # Fit the regression model
    model = sm.OLS(y, X).fit()

    # View the model summary
    summary_text = model.summary().as_text()
    file_path = 'Output/Document/regression_summary_basic_model.txt'
    with open(file_path, 'w') as file:
        file.write(summary_text)
        file.write(summary_with_parameter)

"""







    """df_month_andalucia = imd.import_montly_andalucia()
    df_month_andalucia.columns

    df_month_andalucia = df_month_andalucia.replace(0, np.nan)
    df_month_andalucia = df_month_andalucia.set_index('DATE')
    df_monthcol = df_month.copy()

    gf.plot_correlation_matrix(df_month)


    df_month2 = pd.merge(df_monthcol, df_month_andalucia, left_on='DATE', right_on='DATE', how='left')


    print_doc_descriptive_vars(df_month2,'VIRGEN_EXTRA_EUR_kg')
    """


    def back_testing_regression_OLD(df: pd.DataFrame(), x_cols, y_var, initial_date: str = '2022-05-01',
                                    final_date: str = '2023-05-01', signif: bool = True,
                                    regr_type='Linear', num_variables: int = 4, window: int = 48, step_ahead: int = 12):
        """
        Rolling window hedging. It evaluates the hedging for the selected parameters, it outputs the cash flow for the selected period
        Input:
            df: Dataframe. It takes a df with the objective function, and all spot and forward columns. The hedging is done for all dates in the index.
            x_cols: List of Spot columns. These are the columns that will be selected in the regression
            y_var: String of the name of the objettve fnction (y)
            volumen: List. Volume of SSCC for each month
            initial_date: String. It is the first month in which the hedging is done. It must have the format: 'YYYY-MM-01'
            final_date: String. It is the last month in which the hedging is done. It must have the format: 'YYYY-MM-01'
            signif: Boolean. True to apply the Step Wise method in the regression. False to select all variales passed in x_cols. Only used in the regression.
            prima: Float. Selected adder (Prima). The prima is added in the final calulations as an addition to the base cash flow
            regr_type: String. Linear, for linear regression, or Huber, for robust regression. Only used in the regression.
            num_variables: Integer. Maximum number of variables to select while doing Step Wise regression, it is applied when signif is True. Only used in the regression.
            window: Integer. Training window in which the regression is calibrated
            step_ahead: Integer. Number of months ahead selected to perform hedging in. If step_ahead = 1 it means that the hedging is caculated for 1 month ahead
        Outputs: Dataframe.
            Dataframe with the final hedging calculations for each month, it has the following columns:
                vars: List of the variables used
                coefs: List of the coefficients used in the regression
                real_date: Date (Month) in which the hedging is done (m)
                forward_date: Date (Month) for which the hedging is done
                sscc_estimado: Estimated value of the SSCC of the months in the test dataframe
                sscc_spot_m1: Real value of the SSCC of the months in the test dataframe
                total_liquid: Sum of all liquidations done by the variables
                r2: R2 of the regression. It is the R2 in-sample.
                cash_flow_EUR: It is the Cash Flow resulting of the entire hedging process
                cash_flow_prima_EUR: It is the Cash Flow resulting of the entire hedging process plus an adder
                cash_flow_inicial: It is the Cash Flow resulting of the initial part of the process. It does not take into account the swap liquidations
                cash_flow_EUR_MWh: cash_flow_EUR divided by the volume (volumen)
                cash_flow_prima_EUR_MWh: cash_flow_prima_EUR_MWh divided by the volume (volumen)
                cash_flow_inicial_EUR_MWh: cash_flow_inicial divided by the volume (volumen)
                Cuadrados_Sin_C: It is used for the %Mejora metric, it is cash_flow_inicial_EUR_MWh ^2
                Cuadrados_Con_C: It is used for the %Mejora metric, it is cash_flow_EUR ^2

        """

        df_total = pd.DataFrame()

        # initial_date fecha inicial de prevision
        # final_date fecha final de prevision
        d = df.loc[initial_date:].index[0] - relativedelta(months=window)  # Date defiition
        df = df.loc[:final_date]

        unique_dates = df.index.unique()  # List of dates of the DataFrame
        # unique_dates1 = unique_dates[:-(window) - (step_ahead) + 1]  # row for the rolling windows
        unique_dates1 = unique_dates[:- (step_ahead) + 1]  # List of dates to iterate over
        print(unique_dates1)
        for idx, i in enumerate(unique_dates1):


            ###### Date range interval for the train dataset, delimited by the window size

            date = i.date()  # Starting date of the training window
            date_max = unique_dates[
                idx + window - 1].date()  # End date of the training window. Month in which I perform the hedging

            df_out2 = df.loc[date: date_max]  # Dataframe with training window

            for step in range(1,
                              step_ahead + 1):  # Iteration on each month of the test window delimited by the step_ahead parameter
                # vol_index = step -1
                df_res = pd.DataFrame()

                date_max_step = unique_dates[idx + window + step - 1].date()  # Date of each step

                df_test = df.loc[date_max_step: date_max_step]

                ###### Regression with the forward values

                df_reg = rf.regression_OLD(df_out2, x_cols, y_var, df_test=df_test, reg_type=regr_type,
                                        significativas=signif,
                                        n_vars=num_variables)

                ###### Calculate liquidations: LIQUIDATIONS

                # liquid = []
                vars = df_reg['vars'][0]
                coefs = df_reg['coef'][0]
                mape = df_reg['MAPE'][0]
                # new

                ###### Swap liquidations: SWAPS

                # for numero,c in enumerate(coefs):

                #     factor = c * volumen[vol_index]
                #     if vars[numero] == 'TRAPI2Mc1' or vars[numero] == 'BRT-': # Cotizan en dolares

                #         swap = float(factor * ( df[vars[numero]].loc[date_max_step : date_max_t][0] / df['EUR='].loc[date_max_step : date_max_t][0] -  df_forwards2[vars[numero]].loc[date_max_step : date_max_t][0] / df_forwards2['EUR='].loc[date_max_step : date_max_t][0]   ) ) # swap = factor * ( [spot M +1] - [forward M +1 in M]  )
                #     elif vars[numero] == 'HT':
                #         swap = 0
                #     else:

                #         swap = float(factor *( df[vars[numero]].loc[date_max_step : date_max_t][0] -  df_forwards2[vars[numero]].loc[date_max_step : date_max_t][0]   ) ) # swap = factor * ( [spot M +1] - [forward M +1 in M]  )

                #     liquid.append(swap)

                res_pred = np.concatenate([np.ravel(rr) for rr in df_reg['pred']])

                ###### CALCULATIONS

                df_res['vars'] = [vars]
                df_res['coefs'] = [coefs]
                df_res['mape'] = [mape]
                df_res['real_date'] = date_max
                df_res['forward_date'] = date_max_step
                df_res['sscc_estimado'] = res_pred[0]  # valor predicho
                df_res['sscc_spot_m1'] = float(df_test[y_var][0])  # valor real
                # df_res['total_liquid'] = sum(liquid)
                df_res['r2'] = df_reg['r2'][0]

                # df_res['cash_flow_EUR'] = volumen[vol_index] * (df_res['sscc_estimado'] - df_res['sscc_spot_m1'] ) + df_res['total_liquid']
                # df_res['cash_flow_prima_EUR'] = volumen[vol_index] * (prima + df_res['sscc_estimado'] - df_res['sscc_spot_m1'] ) + df_res['total_liquid']
                # df_res['cash_flow_inicial'] = volumen[vol_index] * (df_res['sscc_estimado'] - df_res['sscc_spot_m1'] )

                # df_res['cash_flow_EUR_MWh'] = df_res['cash_flow_EUR'] / volumen[vol_index]
                # df_res['cash_flow_prima_EUR_MWh'] = df_res['cash_flow_prima_EUR'] / volumen[vol_index]
                # df_res['cash_flow_inicial_EUR_MWh'] = (df_res['sscc_estimado'] - df_res['sscc_spot_m1'] )

                # df_res['Cuadrados_Sin_C'] = df_res['cash_flow_inicial_EUR_MWh']**2
                # df_res['Cuadrados_Con_C'] = df_res['cash_flow_EUR_MWh']**2

                df_total = pd.concat([df_total, df_res], axis=0)

        return df_total.reset_index(drop=True)



























